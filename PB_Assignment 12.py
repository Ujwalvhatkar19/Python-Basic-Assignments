#!/usr/bin/env python
# coding: utf-8

# # Que 1: In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?

# In[ ]:


Answer 1:
PdfFileReader() needs to be opened in read-binary mode by passing 'rb' as the second 
argument to open(). Likewise, the File object passed to PyPDF2. 
PdfFileWriter() needs to be opened in write-binary mode with 'wb'.


# # Que 2: From a PdfFileReader object, how do you get a Page object for page 5?

# In[ ]:


Answer 2: 
pdfreader.getpage()


# # Que 3: What PdfFileReader variable stores the number of pages in the PDF document?

# In[ ]:


Answer 3:
The total number of pages in the document is stored in the numPages attribute 
of a PdfFileReader object 


# # Que 4:If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before you can obtain Page objects from it?

# In[ ]:


Answer 4 :
# import PdfFileWriter and PdfFileReader 
# class from PyPDF2 library
from PyPDF2 import PdfFileWriter, PdfFileReader
  
# Create a PdfFileWriter object
out = PdfFileWriter()
  
# Open encrypted PDF file with the PdfFileReader
file = PdfFileReader("myfile_encrypted.pdf")
  
# Store correct password in a variable password.
password = "swordfish"
  
# Check if the opened file is actually Encrypted
if file.isEncrypted:
  
    # If encrypted, decrypt it with the password
    file.decrypt(swordfish)
  
    # Now, the file has been unlocked. then we can obtain te page object from it. 


# # Que 5:What methods do you use to rotate a page?

# In[ ]:


Answer 5:
target_file = 'filename.pdf'
output_file = open ('filename.pdf' , 'wb')

pdf = PdfFileReader(target_file)
writer = PdfFileWriter()

page = pdf.getpage(0)
page.rotateClockwise(90)
writer.addPage(page)

writer.write(output_file)
output_file.closer()


# # Que 6: What is the difference between a Run object and a Paragraph object?

# In[ ]:


A run is the object most closely associated with inline content; text, pictures, 
and other items that are flowed between the block-item boundaries within a paragraph

Paragraph object providing access to the formatting properties for this paragraph, 
such as line spacing and indentation.


# # Que 7:How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc?

# In[ ]:


Answer 7:
#!pip install python-docx
import docx
doc = docx.Document('abc.docx')
doc.paragraphs
#By using doc.paragraphs 


# # Que 8:What type of object has bold, underline, italic, strike, and outline variables?

# In[ ]:


Answer 8: 
A Run object has bold, underline,italic,strike and outline variables


# # Que 9: What is the difference between False, True, and None for the bold variable?

# In[ ]:


Answer 9: 
Runs can be further styled using text attributes. Each attribute can be 
set to one of three values:
True (the attribute is always enabled, no matter what other styles are applied to the run),
False (the attribute is always disabled),
None (defaults to whatever the run’s style is set to)

True always makes the Run object bolded and False makes it always not bolded, no matter 
what the style’s bold setting is. 
None will make the Run object just use the style’s bold setting


# # Que 10:How do you create a Document object for a new Word document?

# In[ ]:


Answer 10: 
By Calling the docx.Document() function. 


# # Que 11: How do you add a paragraph with the text 'Hello, there!'; to a Document object stored in avariable named doc?

# In[ ]:


Answer 11:
import docx
doc = docx.Document()

doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


# # Que 12:What integers represent the levels of headings available in Word documents?

# In[ ]:


Answer 12:
integer from 0 to 4
The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. 
The integer 0 makes the heading the Title style, which is used for the top of the document. 
Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the 
lowest subheading

